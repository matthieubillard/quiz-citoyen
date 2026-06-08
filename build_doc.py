# -*- coding: utf-8 -*-
import os
import re
import unicodedata
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Import our Q&A modules
import cat_droits_devoirs
import cat_histoire_geo
import cat_principes_valeurs
import cat_systeme_politique
import cat_societe_francaise

# 1. Normalization function to handle special characters, hyphens, and whitespace
def normalize(text):
    text = text.lower()
    # Remove accents
    text = "".join(c for c in unicodedata.normalize('NFD', text) if unicodedata.category(c) != 'Mn')
    # Keep only alphanumeric characters
    text = "".join(c for c in text if c.isalnum())
    return text

# 2. Build the master normalized mapping
master_dict = {}
all_qnas = [
    cat_droits_devoirs.qna,
    cat_histoire_geo.qna,
    cat_principes_valeurs.qna,
    cat_systeme_politique.qna,
    cat_societe_francaise.qna
]

for qna in all_qnas:
    for q, a in qna.items():
        master_dict[normalize(q)] = a

# 3. Helper functions for styling tables in python-docx
def set_cell_background(cell, hex_color):
    """Set the background color of a cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner padding (margins) of a cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, hex_color="D3D3D3"):
    """Set thin clean borders for the table."""
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

# 4. Generate the document
def build_word_document():
    doc = Document()
    
    # Page setup - 2cm margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Charcoal
    
    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("GUIDE DE RÉVISION\nEXAMEN CIVIQUE FRANÇAIS")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x00, 0x33, 0x66) # Deep Blue
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Toutes les questions officielles de connaissances avec leurs réponses rédigées\n(Niveaux : Carte de séjour pluriannuelle, Carte de résident, Naturalisation)")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # Intro Callout Box
    intro_table = doc.add_table(rows=1, cols=1)
    intro_table.autofit = False
    intro_table.columns[0].width = Inches(6.9)
    intro_cell = intro_table.cell(0, 0)
    set_cell_background(intro_cell, "F0F4F8") # Very light blue-gray
    set_cell_margins(intro_cell, top=150, bottom=150, left=200, right=200)
    
    # Left border only for callout style
    tcPr = intro_cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="003366"/>'
        f'  <w:top w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    intro_p = intro_cell.paragraphs[0]
    irun = intro_p.add_run(
        "Ce document a été élaboré pour préparer efficacement l'examen civique obligatoire en France. "
        "Il réunit les questions de connaissances officielles issues du référentiel ministériel. "
        "Pour chaque thématique, les questions sont répertoriées avec les réponses correspondantes pour vous permettre de bachoter dans les meilleures conditions.\n\n"
        "Note : L'examen officiel QCM comporte 40 questions (28 questions de connaissances issues de ce guide et 12 questions de mise en situation confidentielles). "
        "Le seuil de réussite est fixé à 80% (soit 32 bonnes réponses sur 40)."
    )
    irun.font.size = Pt(9.5)
    irun.font.italic = True
    
    doc.add_paragraph() # Spacer
    
    # Read and parse all_questions.txt to generate tables category by category
    current_category = None
    category_questions = []
    
    with open("all_questions.txt", "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    for line in lines:
        line_str = line.strip()
        if line_str.startswith("=== CATEGORY:"):
            # If we already have questions for the previous category, build the section
            if current_category and category_questions:
                build_category_section(doc, current_category, category_questions)
                category_questions = []
            
            # Extract new category name
            m = re.match(r'=== CATEGORY:\s*(.*?)\s*\(\d+\s+questions\)\s*===', line_str)
            if m:
                current_category = m.group(1)
            else:
                current_category = line_str.replace("===", "").strip()
        elif line_str.startswith("- "):
            q = line_str[2:]
            # Filter valid questions
            if q.endswith('?') or q.endswith(':') or '?' in q or 'complétez' in q.lower():
                # Look up answer
                norm_q = normalize(q)
                ans = master_dict.get(norm_q)
                if not ans:
                    # Loose matching fallback in case of trailing spaces or punctuation variations
                    # Let's check if there's a close key
                    for key in master_dict:
                        if key in norm_q or norm_q in key:
                            ans = master_dict[key]
                            break
                if not ans:
                    # In case of absolute fallback
                    print(f"Warning: answer not found for question: {q}")
                    ans = "Réponse à valider."
                category_questions.append((q, ans))
                
    # Build the last category
    if current_category and category_questions:
        build_category_section(doc, current_category, category_questions)
        
    doc.save("guide_revision_examen_civique.docx")
    print("Document successfully created: guide_revision_examen_civique.docx")

def build_category_section(doc, category_name, questions):
    # Page break for clean organization
    if len(doc.paragraphs) > 4: # Don't add page break at the very beginning
        doc.add_page_break()
        
    # Category Heading
    h = doc.add_paragraph()
    h_run = h.add_run(category_name)
    h_run.font.name = 'Arial'
    h_run.font.size = Pt(15)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    # Table Setup
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_borders(table, "D3D3D3")
    
    # Set column widths
    table.columns[0].width = Inches(2.7)
    table.columns[1].width = Inches(4.2)
    
    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].width = Inches(2.7)
    hdr_cells[1].width = Inches(4.2)
    
    set_cell_background(hdr_cells[0], "003366")
    set_cell_background(hdr_cells[1], "003366")
    set_cell_margins(hdr_cells[0], top=120, bottom=120, left=150, right=150)
    set_cell_margins(hdr_cells[1], top=120, bottom=120, left=150, right=150)
    
    # Header Text
    p1 = hdr_cells[0].paragraphs[0]
    r1 = p1.add_run("Question")
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    p2 = hdr_cells[1].paragraphs[0]
    r2 = p2.add_run("Réponse attendue")
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Prevent row split across pages
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:tblHeader'))
    
    # Add Question Rows
    for idx, (q, a) in enumerate(questions):
        row = table.add_row()
        row.cells[0].width = Inches(2.7)
        row.cells[1].width = Inches(4.2)
        
        # Shading for alternating rows
        bg_color = "FFFFFF" if idx % 2 == 0 else "F9FBFD"
        set_cell_background(row.cells[0], bg_color)
        set_cell_background(row.cells[1], bg_color)
        
        set_cell_margins(row.cells[0], top=100, bottom=100, left=150, right=150)
        set_cell_margins(row.cells[1], top=100, bottom=100, left=150, right=150)
        
        # Write Question (Bold)
        pq = row.cells[0].paragraphs[0]
        pq.paragraph_format.space_before = Pt(2)
        pq.paragraph_format.space_after = Pt(2)
        rq = pq.add_run(q)
        rq.font.bold = True
        rq.font.size = Pt(9.5)
        
        # Write Answer
        pa = row.cells[1].paragraphs[0]
        pa.paragraph_format.space_before = Pt(2)
        pa.paragraph_format.space_after = Pt(2)
        ra = pa.add_run(a)
        ra.font.size = Pt(9.5)
        
        # CantSplit option for clean printing
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))

if __name__ == '__main__':
    build_word_document()
